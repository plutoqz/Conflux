"""P3.2 document discovery — D0 directory scan, D1 parsers, D2 rule classifier.

D0: cheap directory/identity discovery — read only path, size, mtime, type;
compute a candidate priority without reading large file bodies.
D1: deterministic parsing — content hash + structured extraction (title,
headings, notebook cells, PDF pages, DOCX paragraphs) with path/line/page
anchors; parse failures stay `partial/failed` without blocking others.
D2: rule classification — decide document kind + authority candidacy from
filename, directory, title and structure.  Rules only produce candidates;
authority confirmation is always a human step (P3 §8.2, §4.5).
"""

from __future__ import annotations

import hashlib
import re
import time
from pathlib import Path
from typing import Any

from ..project_registry.models import ProjectDefinition
from .contracts import (
    ClassificationSource,
    DocumentAuthority,
    DocumentKind,
    ParseStatus,
    ProjectDocument,
)

# Candidate directories scanned for research documents (P3 §8.1).
DOC_DIRS: tuple[str, ...] = (
    "docs", "notes", "experiments", "reports", "papers", "plan", "plans",
    "research", "deliverables", "data/notes", "archive",
)

# Root-level convention documents (P3 §8.1).
ROOT_DOC_PREFIXES: tuple[str, ...] = (
    "PROJECT", "README", "AGENTS", "DESIGN", "PRODUCT", "PLAN", "ROADMAP",
    "CHANGELOG", "CONTRACT", "TODO", "NOTES",
)

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({
    ".md", ".markdown", ".txt", ".pdf", ".docx", ".ipynb", ".yaml", ".yml",
    ".json", ".csv", ".tsv", ".rst", ".org",
})

IGNORED_PARTS: frozenset[str] = frozenset({
    ".git", ".venv", "venv", "node_modules", "__pycache__", ".pytest_cache",
    ".mypy_cache", ".ruff_cache", "dist", "build", ".tox", ".idea", ".vscode",
    "site-packages", ".tmp",
})

DEFAULT_IGNORE_PATTERNS: tuple[str, ...] = (
    "*.pyc", "*.pyo", "*.exe", "*.dll", "*.so", "*.dylib", "*.png", "*.jpg",
    "*.jpeg", "*.gif", "*.svg", "*.ico", "*.woff", "*.woff2", "*.ttf",
    "*.zip", "*.tar", "*.gz", "*.7z", "*.rar", "*.sqlite", "*.db", "*.lock",
    "*.min.js", "*.min.css", "*.map",
)

MAX_FILE_BYTES = 4 * 1024 * 1024       # skip bodies above 4 MB (D0-only)
MAX_CANDIDATES_PER_SCAN = 2000


def document_id_for(project_id: str, relative_path: str) -> str:
    """Stable id: sha256(project_id + '|' + relative_path)[:16]."""
    raw = f"{project_id}|{relative_path}".encode("utf-8")
    return hashlib.sha256(raw).hexdigest()[:16]


# ── D0: directory & identity discovery ──────────────────────────


def _ignored(path: Path, ignore_patterns: tuple[str, ...]) -> bool:
    import fnmatch

    name = path.name
    if name in IGNORED_PARTS:
        return True
    if path.is_dir():
        return False
    for pattern in ignore_patterns:
        if fnmatch.fnmatch(name, pattern):
            return True
    return False


def discover_candidates(
    project: ProjectDefinition,
    *,
    ignore_patterns: tuple[str, ...] = DEFAULT_IGNORE_PATTERNS,
    max_files: int = MAX_CANDIDATES_PER_SCAN,
) -> list[dict[str, Any]]:
    """D0: walk the project root, collect candidate document identities.

    Only reads stat() metadata — no file bodies.  Returns dicts with
    path/relative_path/size/mtime/ext and a candidate priority.
    """
    root = Path(project.path).expanduser().resolve()
    candidates: list[dict[str, Any]] = []
    if not root.is_dir():
        return candidates

    base_depth = len(root.parts)

    def _walk(directory: Path) -> None:
        if len(candidates) >= max_files:
            return
        try:
            entries = sorted(directory.iterdir(), key=lambda p: p.name.casefold())
        except OSError:
            return
        for entry in entries:
            if _ignored(entry, ignore_patterns):
                continue
            if entry.is_dir():
                _walk(entry)
                continue
            ext = entry.suffix.casefold()
            if ext not in SUPPORTED_EXTENSIONS:
                continue
            try:
                size = entry.stat().st_size
            except OSError:
                continue
            if size > MAX_FILE_BYTES:
                continue
            relative = entry.relative_to(root).as_posix()
            depth = len(entry.parts) - base_depth
            candidates.append({
                "path": str(entry),
                "relative_path": relative,
                "size": size,
                "mtime": entry.stat().st_mtime,
                "ext": ext,
                "depth": depth,
                "priority": _candidate_priority(relative, depth),
            })

    _walk(root)
    candidates.sort(key=lambda c: (-c["priority"], c["relative_path"]))
    return candidates[:max_files]


def _candidate_priority(relative_path: str, depth: int) -> int:
    """Higher = more likely a research document (P3 §8.2 D0)."""
    lowered = relative_path.casefold()
    score = 0
    first_part = lowered.split("/")[0] if "/" in lowered else ""
    if first_part in ("docs", "notes", "papers", "plan", "plans", "research", "reports", "experiments"):
        score += 3
    name = Path(lowered).name
    stem = Path(lowered).stem.upper()
    if any(name.upper().startswith(prefix) for prefix in ROOT_DOC_PREFIXES):
        score += 4
    # Root-level convention docs (README/PROJECT/AGENTS/DESIGN/PRODUCT) at
    # depth 1 outrank everything else (P3 §8.1).
    if depth == 1 and stem in ("README", "PROJECT", "AGENTS", "DESIGN", "PRODUCT", "PLAN"):
        score += 6
    if depth <= 2:
        score += 1
    if name.endswith((".md", ".markdown", ".txt", ".org", ".rst")):
        score += 2
    return score


# ── D1: deterministic parsing ───────────────────────────────────


def parse_document(
    path: Path,
    *,
    content: bytes | None = None,
    extractor_version: str = "d1-20260812",
) -> ProjectDocument:
    """Parse one file into a ProjectDocument (D1).

    Reads the file once, hashes it, extracts title + headings/anchors by
    extension.  Unsupported/oversized files return parse_status=unsupported.
    """
    try:
        if content is None:
            content = path.read_bytes()
    except OSError:
        return _failed_document(path, ParseStatus.FAILED)
    content_hash = hashlib.sha256(content).hexdigest()[:32]
    ext = path.suffix.casefold()

    text, anchors, title, meta = _extract_text(path, content, ext)

    if text is None:
        status = ParseStatus.UNSUPPORTED
        title = ""
        text = ""
    elif not text.strip():
        status = ParseStatus.READY if ext in (".yaml", ".yml", ".json") else ParseStatus.PARTIAL
    else:
        status = ParseStatus.READY

    return ProjectDocument(
        document_id=document_id_for("", path.name),  # placeholder; caller sets project
        project_id="",
        path=str(path),
        content_hash=content_hash,
        kind=DocumentKind.OTHER,
        authority=DocumentAuthority.CANDIDATE,
        parse_status=status,
        language=_detect_language(text[:4000]),
        title=title,
        modified_at=path.stat().st_mtime if path.exists() else 0.0,
        indexed_at=time.time(),
        extractor_version=extractor_version,
        classification_source=ClassificationSource.RULE,
        classification_confidence=0.0,
        metadata={
            "size": len(content),
            "ext": ext,
            "headings": anchors[:50],
            "text_len": len(text),
        },
    )


def _failed_document(path: Path, status: ParseStatus) -> ProjectDocument:
    return ProjectDocument(
        document_id=document_id_for("", path.name),
        project_id="",
        path=str(path),
        parse_status=status,
        extractor_version="d1-20260812",
    )


def _extract_text(
    path: Path,
    content: bytes,
    ext: str,
) -> tuple[str | None, list[str], str, dict[str, Any]]:
    """Return (text, heading anchors, title, metadata); text=None if unsupported."""
    if ext in (".md", ".markdown", ".txt", ".rst", ".org"):
        return _parse_markdown(content)
    if ext == ".ipynb":
        return _parse_notebook(content)
    if ext == ".pdf":
        return _parse_pdf(path, content)
    if ext == ".docx":
        return _parse_docx(path)
    if ext in (".yaml", ".yml", ".json"):
        return (content.decode("utf-8", errors="replace"), [], "", {})
    if ext in (".csv", ".tsv"):
        return (content.decode("utf-8", errors="replace")[:100_000], [], "", {})
    return (None, [], "", {})


def _parse_markdown(content: bytes) -> tuple[str, list[str], str, dict[str, Any]]:
    text = content.decode("utf-8", errors="replace")
    headings: list[str] = []
    for line in text.splitlines()[:400]:
        stripped = line.strip()
        if stripped.startswith("#"):
            headings.append(stripped[:200])
        elif re.match(r"^(=+|-+)\s*$", stripped) and headings:
            headings.append(headings[-1])  # underline-style heading
    title = ""
    for line in text.splitlines()[:40]:
        stripped = line.strip().strip("#").strip()
        if stripped and not stripped.startswith(("```", "<!--", "[", ">")):
            title = stripped[:200]
            break
    return text, headings, title, {"format": "markdown"}


def _parse_notebook(content: bytes) -> tuple[str, list[str], str, dict[str, Any]]:
    import json as _json

    try:
        notebook = _json.loads(content.decode("utf-8", errors="replace"))
    except ValueError:
        return "", [], "", {"format": "ipynb", "error": "invalid json"}
    cells = notebook.get("cells", []) if isinstance(notebook, dict) else []
    parts: list[str] = []
    headings: list[str] = []
    for cell in cells:
        cell_type = cell.get("cell_type", "")
        source = "".join(cell.get("source", []))
        if cell_type == "markdown":
            for line in source.splitlines():
                stripped = line.strip()
                if stripped.startswith("#"):
                    headings.append(stripped[:200])
            parts.append(source)
        elif cell_type == "code":
            parts.append(f"```python\n{source}\n```")
    title = notebook.get("metadata", {}).get("title", "") if isinstance(notebook, dict) else ""
    return "\n\n".join(parts), headings, title, {"format": "ipynb"}


def _parse_pdf(path: Path, content: bytes) -> tuple[str, list[str], str, dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return (None, [], "", {"format": "pdf", "error": "pypdf missing"})
    try:
        import io as _io

        reader = PdfReader(_io.BytesIO(content))
        pages: list[str] = []
        page_anchors: list[str] = []
        for index, page in enumerate(reader.pages[:120]):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""
            pages.append(page_text)
            first_lines = [l for l in page_text.splitlines() if l.strip()][:3]
            if first_lines:
                page_anchors.append(f"p{index + 1}:{first_lines[0][:120]}")
        text = "\n\n".join(pages)
        title = ""
        if text.strip():
            first = [l for l in text.splitlines() if l.strip()]
            title = first[0][:200] if first else ""
        return text, page_anchors, title, {"format": "pdf", "pages": len(reader.pages)}
    except Exception as exc:
        return "", [], "", {"format": "pdf", "error": str(exc)[:120]}


def _parse_docx(path: Path) -> tuple[str, list[str], str, dict[str, Any]]:
    try:
        from docx import Document as _Document
    except ImportError:
        return (None, [], "", {"format": "docx", "error": "python-docx missing"})
    try:
        document = _Document(str(path))
        parts: list[str] = []
        headings: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = str(paragraph.style.name or "").casefold()
            if "heading" in style or "标题" in style:
                headings.append(text[:200])
            parts.append(text)
        title = parts[0][:200] if parts else ""
        return "\n".join(parts), headings, title, {"format": "docx", "paragraphs": len(parts)}
    except Exception as exc:
        return "", [], "", {"format": "docx", "error": str(exc)[:120]}


def _detect_language(text: str) -> str:
    if not text.strip():
        return ""
    sample = text[:2000]
    cjk = sum(1 for ch in sample if "\u4e00" <= ch <= "\u9fff")
    # Short docs: any CJK characters in the first line usually mean Chinese.
    first_line = sample.splitlines()[0] if sample.splitlines() else sample
    if cjk >= 3 or any("\u4e00" <= ch <= "\u9fff" for ch in first_line):
        return "zh"
    return "en"


# ── D2: rule classification ─────────────────────────────────────


KIND_HINTS: list[tuple[tuple[str, ...], DocumentKind]] = [
    (("charter", "纲领", "project.md", "proposal"), DocumentKind.CHARTER),
    (("plan", "计划", "roadmap", "road_map", "路线图", "milestone", "里程碑", "sprint"), DocumentKind.PLAN),
    (("decision", "决策", "adr-", "adr_", "结论"), DocumentKind.DECISION),
    (("experiment", "实验", "benchmark", "评测", "ablation", "消融"), DocumentKind.EXPERIMENT),
    (("report", "报告", "deliverable", "交付", "summary", "总结"), DocumentKind.REPORT),
    (("paper", "论文", "paper_note", "文献", "reading", "notes/"), DocumentKind.PAPER_NOTE),
    (("design", "设计", "architecture", "架构", "api", "spec"), DocumentKind.CODE_DOC),
]


def classify_document(document: ProjectDocument) -> ProjectDocument:
    """D2: rule-based kind classification + candidate priority (no authority)."""
    lowered_path = str(document.path).casefold()
    lowered_title = (document.title or "").casefold()

    best_kind = DocumentKind.OTHER
    best_score = 0
    for hints, kind in KIND_HINTS:
        score = 0
        for hint in hints:
            if hint in lowered_path:
                score += 2
            if hint in lowered_title:
                score += 1
        if score > best_score:
            best_score = score
            best_kind = kind

    document.kind = best_kind
    document.classification_source = ClassificationSource.RULE
    document.classification_confidence = round(min(1.0, best_score / 4.0), 2)
    # Authority is NEVER set by rules (P3 §4.5).
    document.authority = DocumentAuthority.CANDIDATE
    return document
